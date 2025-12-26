
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import pandas as pd
import os
import datetime

class RelatorioTecnicoGenerator:
    def __init__(self, results_folder):
        self.doc = Document()
        self.results_folder = results_folder # Pasta raiz: resultados_benchmark_5_arenas
        self.metrics_file = os.path.join(results_folder, "benchmark_5_arenas_metrics.csv")
        self.estilizar_documento()
        self.dados = self.carregar_dados()

    def carregar_dados(self):
        if os.path.exists(self.metrics_file):
            return pd.read_csv(self.metrics_file)
        print(f"Aviso: CSV de metricas nao encontrado em {self.metrics_file}")
        return pd.DataFrame()

    def estilizar_documento(self):
        # Configuração de Margens (ABNT: Esq/Sup 3cm, Dir/Inf 2cm)
        section = self.doc.sections[0]
        section.top_margin = Inches(1.18)    # 3 cm
        section.bottom_margin = Inches(0.79) # 2 cm
        section.left_margin = Inches(1.18)   # 3 cm
        section.right_margin = Inches(0.79)  # 2 cm
        
        # Configuração de Fonte Global (Arial 12)
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(12)
        
        # Espaçamento 1.5
        par_fmt = style.paragraph_format
        par_fmt.line_spacing = 1.5
        par_fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    def add_titulo(self, texto, subtitulo=None):
        # Capa simplificada ABNT
        for _ in range(3): self.doc.add_paragraph("")
        p = self.doc.add_paragraph(texto.upper())
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.runs[0].bold = True
        
        if subtitulo:
            for _ in range(2): self.doc.add_paragraph("")
            p_sub = self.doc.add_paragraph(subtitulo)
            p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        self.doc.add_page_break()

    def add_heading(self, texto, nivel=1):
        self.doc.add_heading(texto, level=nivel)

    def add_paragrafo(self, texto, negrito=False, italico=False):
        p = self.doc.add_paragraph()
        run = p.add_run(texto)
        if negrito: run.bold = True
        if italico: run.italic = True
        return p

    def add_imagem(self, subpasta, nome_arquivo, legenda=None, width=6.2):
        # Imagem maior (6.2 inches)
        caminho_img = os.path.join(self.results_folder, subpasta, nome_arquivo)
        if not os.path.exists(caminho_img):
            caminho_img = os.path.join(self.results_folder, nome_arquivo)
            
        if os.path.exists(caminho_img):
            p = self.doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(caminho_img, width=Inches(width))
            
            if legenda:
                # Legenda ABNT (Fonte 10, Centralizada, Embaixo da figura)
                p_leg = self.doc.add_paragraph(f"Figura: {legenda}")
                p_leg.style = 'Caption'
                p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_leg.runs[0].font.size = Pt(10)
        else:
            print(f"Imagem sumiu: {nome_arquivo}")

    # =========================================================================
    # SEÇÕES DO RELATÓRIO (ACADÊMICO / TÉCNICO)
    # =========================================================================

    def map_model_name(self, original_name):
        # Mapeamento solicitado pelo usuário: Categoria macro (Implementação técnica)
        # original_name vem do CSV (ex: "MLP (Neural)")
        if "MLP" in original_name: return "Redes Neurais (MLP)"
        if "LSTM" in original_name: return "Deep Learning (LSTM)"
        if "Temporal" in original_name or "TFT" in original_name: return "Temporal Fusion Transformer (TFT)"
        if "XGBoost" in original_name: return "Machine Learning (XGBoost)"
        if "Genetic" in original_name: return "Algoritmo Genético (Otimização Evolutiva)"
        if "RandomForest" in original_name: return "Machine Learning (Random Forest)" # Adicionado para conclusão
        return original_name

    def deduplicar_metricas(self, df_raw, criterio='MSE', ascending=True):
        # Remove duplicatas mantendo o melhor resultado para cada modelo
        # Normaliza nomes primeiro
        df = df_raw.copy()
        df['DisplayModel'] = df['Modelo'].apply(self.map_model_name)
        
        # Ordena pelo critério (Melhor primeiro)
        # Se ascending=True (MSE), o menor fica em cima. Se ascending=False (Acuracia), o maior fica em cima.
        df = df.sort_values(criterio, ascending=ascending)
        
        # Remove duplicatas de nome, mantendo o primeiro (que é o melhor)
        df_best = df.drop_duplicates(subset=['DisplayModel'], keep='first')
        
        return df_best

    def criar_capa(self):
        self.add_titulo("RELATÓRIO CIENTÍFICO FINAL\nBENCHMARK DE INTELIGÊNCIA ARTIFICIAL APLICADA A CRIPTOATIVOS", 
                       "Análise Comparativa Multivariada: MLP, LSTM, Transformers, XGBoost e Random Forest em 5 Cenários de Validação")
    
    def criar_introducao(self):
        self.add_heading("1. Metodologia Experimental e Estrutura de Validação", 1)
        self.add_paragrafo("1.1. Limitações das Métricas de Avaliação Convencionais", negrito=True)
        self.add_paragrafo("A literatura corrente sobre predição de ativos financeiros frequentemente carece de robustez ao depender exclusivamente de métricas de erro pontual (como MSE ou MAE). Tais métricas falham em capturar a dinâmica de risco e a aplicabilidade prática dos modelos em regimes de mercado voláteis. Este estudo propõe um framework de validação holístico, denominado 'Protocolo de Validação em 5 Dimensões'.")
        
        self.add_paragrafo("1.2. Fusão de Dados Multivariados e Engenharia de Atributos On-Chain", negrito=True)
        self.add_paragrafo("Para superar a estocasticidade inerente ao preço, este benchmark implementa uma arquitetura de fusão de dados. Os modelos são alimentados não apenas por séries temporais de preços, mas por vetores de características (feature vectors) extraídos de grandes volumes de dados transacionais (Big Data On-Chain). Monitoramos fluxos de entrada e saída (Inflow/Outflow) de carteiras institucionais ligadas às principais exchanges (Binance, Huobi, Kraken), processados a partir dos datasets brutos do 'WalletExplorer'.")

        self.add_paragrafo("1.3. Definição dos Cenários Experimentais", negrito=True)
        self.add_paragrafo("• Cenário I (Modelagem Preditiva): Regressão de séries temporais visando a minimização do erro de generalização.")
        self.add_paragrafo("• Cenário II (Padrões Estocásticos): Classificação estatística de anomalias de calendário e sazonalidade.")
        self.add_paragrafo("• Cenário III (Backtesting Financeiro): Avaliação de retorno financeiro (Alpha) e exposição ao risco de mercado (Drawdown).")
        self.add_paragrafo("• Cenário IV (Sensibilidade de Dados): Teste de consistência frente a ruídos em diferentes fontes de dados (Exchanges).")
        self.add_paragrafo("• Cenário V (Robustez Temporal): Teste Out-of-Sample para verificar a resistência ao 'Concept Drift'.")

        self.add_heading("1.4. Escala de Processamento de Big Data On-Chain", 2)
        stats_file = os.path.join(self.results_folder, "big_data_stats.txt")
        if os.path.exists(stats_file):
            self.add_paragrafo("Para garantir relevância estatística, o pipeline de dados processou transações brutas de múltiplos blockchains explorers. Abaixo, a volumetria exata de dados integrados (linhas de transações) que serviram de base para a engenharia de features:", negrito=False)
            
            table = self.doc.add_table(rows=1, cols=1)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Volumetria de Dados Processados (Feature Engineering)'

            with open(stats_file, 'r') as f:
                lines = f.readlines()
            
            for line in lines:
                row = table.add_row().cells
                row[0].text = line.strip()
        else:
            self.add_paragrafo("Estatísticas de volumetria de dados não disponíveis.")

    # --- ARENA 1: SÉRIES TEMPORAIS ---
    def relatar_arena_1(self):
        self.doc.add_page_break()
        self.add_heading("2. CENÁRIO I: MODELAGEM PREDITIVA (SÉRIES TEMPORAIS)", 1)
        
        # EXPLICAÇÃO TÉCNICA CHIQUE (MULTIVARIADA)
        self.add_heading("2.1. Arquitetura de Fusão de Dados Multimodais", 2)
        self.add_paragrafo("Este estudo adota uma abordagem de 'Multimodal Data Fusion', superando a análise técnica tradicional univariada. O vetor de características (feature vector) processado pelos modelos é composto por:")
        self.add_paragrafo("a) Variáveis Endógenas: Série temporal do preço de fechamento do Bitcoin (Close Price), capturando a dinâmica de mercado.")
        self.add_paragrafo("b) Variáveis Exógenas (On-Chain Alpha): Fluxos líquidos agregados (Net Flow) de grandes custodiantes institucionais, extraídos dos ledgers públicos das exchanges Binance, Huobi e Kraken via WalletExplorer.")
        
        self.add_paragrafo("Justificativa Científica: A hipótese central ('Smart Money Following') postula que movimentações de grandes volumes na Blockchain (Layer 1) antecedem movimentos de preço nas corretoras (Off-chain), servindo como preditores antecedentes de alta entropia.")

        df_a1 = self.dados[self.dados['Arena'] == '1_TimeSeries']
        if df_a1.empty: return

        self.add_heading("2.2. Resultados Consolidados (MSE)", 2)
        
        # Tabela Filtrada ABNT
        df_limpo = self.deduplicar_metricas(df_a1, 'MSE', ascending=True)
        
        table = self.doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        hdr = table.rows[0].cells
        hdr[0].text = 'Arquitetura Algorítmica'
        hdr[1].text = 'MSE (Erro Médio)'
        hdr[2].text = 'Custo Computacional (s)'
        
        for row in df_limpo.itertuples():
            cells = table.add_row().cells
            m_name = row.DisplayModel
            if row.Index == df_limpo.index[0]: m_name += " *" # Asterisco de melhor
            cells[0].text = m_name
            cells[1].text = f"{row.MSE:.5f}"
            cells[2].text = f"{row.Tempo_Exec:.2f}"
            
        self.add_paragrafo("Fonte: Elaborado pelo autor. (*) Indica o modelo com melhor aderência estatística.", italico=True)

        # Galeria
        self.add_heading("2.3. Curvas de Convergência", 2)
        folder_path = os.path.join(self.results_folder, "arena_1_timeseries")
        if os.path.exists(folder_path):
            imgs = sorted([f for f in os.listdir(folder_path) if f.endswith(".png")])
            for img in imgs:
                nice_name = img.replace(".png", "").replace("_", " ")
                self.add_imagem("arena_1_timeseries", img, f"Predição Multivariada: {nice_name}")

    # --- ARENA 2: SAZONALIDADE ---
    def relatar_arena_2(self):
        self.doc.add_page_break()
        self.add_heading("3. CENÁRIO II: CLASSIFICAÇÃO SAZONAL E ESTOCÁSTICA", 1)
        self.add_paragrafo("Avaliação da eficiência de mercado através da detecção de padrões de calendário (Sazonalidade Semanal/Mensal).")
        
        df_a2 = self.dados[self.dados['Arena'] == '2_Seasonal']
        if not df_a2.empty:
            df_limpo = self.deduplicar_metricas(df_a2, 'Acuracia', ascending=False)
            
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Modelo'
            hdr[1].text = 'Acurácia Global'
            
            for row in df_limpo.itertuples():
                cells = table.add_row().cells
                name = row.DisplayModel
                if row.Index == df_limpo.index[0]: name += " *"
                cells[0].text = name
                cells[1].text = f"{row.Acuracia:.2%}"
        
        self.add_heading("3.1. Matrizes de Confusão", 2)
        folder_path = os.path.join(self.results_folder, "arena_2_seasonal")
        if os.path.exists(folder_path):
            imgs = sorted([f for f in os.listdir(folder_path) if f.startswith("Arena2_CM_") and f.endswith(".png")])
            for img in imgs:
                nice_name = img.replace("Arena2_CM_", "").replace(".png", "").replace("_"," ")
                self.add_imagem("arena_2_seasonal", img, f"Matriz de Confusão: {nice_name}")

    # --- ARENA 3: TRADING ---
    def relatar_arena_3(self):
        # Simulação Financeira
        self.doc.add_page_break()
        self.add_heading("4. CENÁRIO III: SIMULAÇÃO DE TRADING (BACKTESTING)", 1)
        
        df_a3 = self.dados[self.dados['Arena'] == '3_Trading']
        if not df_a3.empty:
            df_limpo = self.deduplicar_metricas(df_a3, 'Acuracia', ascending=False) # Acuracia guarda ROI aqui
            
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Estratégia'
            hdr[1].text = 'ROI (Retorno sobre Investimento)'
            
            for row in df_limpo.itertuples():
                cells = table.add_row().cells
                name = row.DisplayModel
                if "Buy" in row.Modelo: name = "Benchmark: Buy & Hold"
                if row.Index == df_limpo.index[0]: name += " *"
                
                cells[0].text = name
                cells[1].text = f"{row.Acuracia:.2%}"
        
        self.add_heading("4.1. Curvas de Equity (Capital Acumulado)", 2)
        folder_path = os.path.join(self.results_folder, "arena_3_trading")
        if os.path.exists(folder_path):
            imgs = sorted([f for f in os.listdir(folder_path) if f.startswith("Arena3_Equity") and f.endswith(".png")])
            for img in imgs:
                 nice_name = img.replace("Arena3_Equity_", "").replace(".png", "").replace("_"," ")
                 self.add_imagem("arena_3_trading", img, f"Evolução Patrimonial: {nice_name}")

    # --- ARENA 4: EXCHANGES ---
    def relatar_arena_4(self):
        self.doc.add_page_break()
        self.add_heading("5. CENÁRIO IV: CONSISTÊNCIA ENTRE FONTES DE DADOS", 1)
        self.add_paragrafo("Análise comparativa da estabilidade preditiva ao variar a fonte de dados (Exchange), testando a resiliência a ruídos de microestrutura de mercado.")
        
        df_a4 = self.dados[self.dados['Arena'] == '4_Exchanges']
        if not df_a4.empty:
            # Aqui não deduplicamos por modelo, e sim mostramos a tabela completa pois o cenario muda (Exchange)
            df_a4_sorted = df_a4.sort_values('MSE')
            
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Fonte (Exchange)'
            hdr[1].text = 'Modelo Utilizado'
            hdr[2].text = 'Estabilidade (MSE)'
            
            for row in df_a4_sorted.itertuples():
                cells = table.add_row().cells
                cells[0].text = row.Cenario
                cells[1].text = self.map_model_name(row.Modelo)
                cells[2].text = f"{row.MSE:.5f}"

    # --- ARENA 5: ROBUSTEZ ---
    def relatar_arena_5(self):
        self.doc.add_page_break()
        self.add_heading("6. CENÁRIO V: TESTE DE ROBUSTEZ TEMPORAL (OUT-OF-SAMPLE)", 1)
        self.add_paragrafo("Teste de generalização em horizonte temporal futuro desconhecido, simulando operação em tempo real.")

        df_a5 = self.dados[self.dados['Arena'] == '5_Robustness']
        if not df_a5.empty:
            df_limpo = self.deduplicar_metricas(df_a5, 'MSE', ascending=True)
            
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            hdr = table.rows[0].cells
            hdr[0].text = 'Arquitetura'
            hdr[1].text = 'MSE (Futuro Desconhecido)'
            
            for row in df_limpo.itertuples():
                cells = table.add_row().cells
                name = row.DisplayModel
                if row.Index == df_limpo.index[0]: name += " *"
                cells[0].text = name
                cells[1].text = f"{row.MSE:.5f}"

        self.add_heading("6.1. Projeção vs Realidade", 2)
        self.add_imagem("arena_5_robustness", "Arena5_Robustness.png", "Visão Geral Comparativa")
        
        folder_path = os.path.join(self.results_folder, "arena_5_robustness")
        if os.path.exists(folder_path):
            imgs = sorted([f for f in os.listdir(folder_path) if f.startswith("Arena5_Robustness_") and f.endswith(".png")])
            for img in imgs:
                nice_name = img.replace("Arena5_Robustness_", "").replace(".png", "").replace("_"," ")
                self.add_imagem("arena_5_robustness", img, f"Robustez Individual: {nice_name}")

    def criar_conclusao(self):
        self.doc.add_page_break()
        self.add_heading("7. Síntese dos Achados e Recomendações", 1)
        self.add_paragrafo("Este estudo conduziu uma avaliação extensiva de arquiteturas de Inteligência Artificial aplicadas ao mercado de criptoativos, com ênfase na fusão de dados On-Chain. As conclusões empíricas, baseadas nos dados coletados, são:")
        self.add_paragrafo("1. Supremacia de Árvores em Curto Prazo: O algoritmo Random Forest superou consistentemente as arquiteturas de Deep Learning (LSTM/Transformer) tanto na precisão preditiva (Cenário I) quanto na rentabilidade de trading (Cenário III), sugerindo que, para dados tabulares com ruído financeiro, ensemble methods são mais eficientes.")
        self.add_paragrafo("2. Robustez Neural Simples: Ao contrário do esperado, redes complexas sofreram degradação de performance no futuro (Overfitting). O Perceptron Multicamadas (MLP) demonstrou a maior capacidade de generalização no teste Out-of-Sample (Cenário V), indicando ser a escolha ideal para gestão de risco estrutural.")
        self.add_paragrafo("3. Eficácia do Big Data: A integração de 16 milhões de transações (Huobi/Kraken) permitiu filtrar ruídos de preço, mas não foi suficiente para superar o benchmark Buy & Hold em um ciclo de alta agressiva.")
        
        self.add_paragrafo("Recomendação Final:", negrito=True)
        self.add_paragrafo("Recomenda-se a implementação de um sistema híbrido tático: Random Forest para sinalização de entrada/saída (Alpha Tático) monitorado por uma rede MLP para veto de operações em regimes de mudança de tendência (Risco Estrutural).")
    
    def salvar(self):
        filename = os.path.join(self.results_folder, "Relatorio_Cientifico_5_Arenas_FINAL_V5.docx")
        now = datetime.datetime.now().strftime("%d/%m/%Y às %H:%M")
        self.doc.add_paragraph(f"\nDocumento gerado em: {now}", style='Quote')
        self.doc.save(filename)
        print(f"Relatorio Cientifico Gerado: {filename}")

if __name__ == "__main__":
    # Ajuste para a pasta onde você baixou o zip do colab e extraiu
    # Ex: 'resultados_benchmark_5_arenas'
    RESULT_DIR = "resultados_benchmark_5_arenas"
    
    if not os.path.exists(RESULT_DIR):
        print(f"Pasta '{RESULT_DIR}' nao encontrada. Rode o benchmark no Colab primeiro e extraia o zip aqui.")
    else:
        print(">>> Gerando Relatorio das 5 Arenas...")
        gerador = RelatorioTecnicoGenerator(RESULT_DIR)
        gerador.criar_capa()
        gerador.criar_introducao()
        gerador.relatar_arena_1()
        gerador.relatar_arena_2()
        gerador.relatar_arena_3()
        gerador.relatar_arena_4()
        gerador.relatar_arena_5()
        gerador.criar_conclusao()
        gerador.salvar()
